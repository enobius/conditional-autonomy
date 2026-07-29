from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('outputs/Thesis_Architecture_Formalization_v1.1.2.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '2E5E7E'; DARK = '17324D'; LIGHT = 'EAF1F5'; PALE = 'F5F7F9'; GRAY = '5B6570'; RED = '8B2F2F'; GOLD = '8A6A16'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None: tblInd = OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in('w:tcW')
            if tcW is None: tcW=OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'),str(widths[i])); tcW.set(qn('w:type'),'dxa')
            set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); trPr.append(hdr)

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'
    for i,h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i], LIGHT)
        p=t.rows[0].cells[i].paragraphs[0]; r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK); r.font.size=Pt(9)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            p=cells[i].paragraphs[0]; r=p.add_run(str(val)); r.font.size=Pt(9)
    set_table_widths(t,widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run('Page '); run.font.size=Pt(9); run.font.color.rgb=RGBColor.from_string(GRAY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)

def add_bullet(doc,text,level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p

def add_number(doc,text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p

def add_code(doc,text):
    p=doc.add_paragraph(style='Code Block'); p.add_run(text); return p

def callout(doc,label,text,color=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(t,[9360]); set_cell_shading(t.cell(0,0), PALE)
    set_repeat_table_header(t.rows[0])
    p=t.cell(0,0).paragraphs[0]; r=p.add_run(label.upper()+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.49)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string('20262D')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name,size,color,before,after in [('Title',28,DARK,0,8),('Subtitle',13,GRAY,0,18),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',11,DARK,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name!='Subtitle'
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for n in ['List Bullet','List Bullet 2','List Number']:
    styles[n].font.name='Calibri'; styles[n].font.size=Pt(11); styles[n].paragraph_format.space_after=Pt(4); styles[n].paragraph_format.line_spacing=1.10
code=styles.add_style('Code Block',WD_STYLE_TYPE.PARAGRAPH); code.font.name='Consolas'; code.font.size=Pt(8.5); code.font.color.rgb=RGBColor.from_string(DARK)
code.paragraph_format.left_indent=Inches(.25); code.paragraph_format.right_indent=Inches(.15); code.paragraph_format.space_before=Pt(3); code.paragraph_format.space_after=Pt(6)

header=sec.header.paragraphs[0]; header.text='THESIS ARCHITECTURE  |  FORMAL RESEARCH SPECIFICATION'; header.runs[0].font.size=Pt(8); header.runs[0].font.bold=True; header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
add_page_number(sec.footer.paragraphs[0])

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(80); r=p.add_run('RESEARCH SPECIFICATION'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(BLUE)
doc.add_paragraph('Conditional Autonomy for Self-Improving Small-Model Agents',style='Title')
doc.add_paragraph('A formal architecture, learning protocol, and REALM-Bench evaluation plan',style='Subtitle')
add_table(doc,['Field','Specification'],[
    ('Working thesis','Small models can achieve reliable agentic planning through typed environments, constrained execution, trace-driven post-training, and modular specialization.'),
    ('Architecture','Supervisor-worker control plane with deterministic capability and validation layers.'),
    ('Learning model','Versioned offline improvement cycles with independent evaluation, promotion gates, rollback, and adapter registry.'),
    ('Primary benchmark','REALM-Bench planning, scheduling, coordination, and disruption tasks.'),
    ('Document status','Version 1.1.2 - converged specification and implementation contract'),
], [1900,7460])
callout(doc,'Core claim','The research target is not unrestricted autonomy. It is measurable conditional autonomy: act when permissions and constraints are satisfied, repair when the plan is recoverable, ask when information is insufficient, and escalate when residual risk exceeds policy.')
doc.add_page_break()

doc.add_heading('1. Research Contribution',level=1)
doc.add_heading('1.1 Primary research question',level=2)
doc.add_paragraph('Can small language models, embedded in a constrained supervisor-worker architecture, improve planning and recovery from execution traces while preserving hard safety invariants and avoiding regressions on previously mastered tasks?')
doc.add_heading('1.2 Proposed contribution',level=2)
add_bullet(doc,'A typed agent-environment interface that separates observed state, canonical state, action proposals, execution permissions, and state transitions.')
add_bullet(doc,'A hybrid control architecture combining iterative agent loops with deterministic graph boundaries, capability enforcement, and constraint validation.')
add_bullet(doc,'A trace-to-training pipeline that turns verified outcomes and failures into SFT, preference, process-reward, or agent-RL data.')
add_bullet(doc,'A modular skill layer using versioned LoRA adapters with explicit routing, compatibility checks, fallbacks, and regression tests.')
add_bullet(doc,'An evaluation protocol that isolates gains from orchestration, validation, post-training, specialization, routing, and repeated improvement cycles.')
doc.add_heading('1.3 Non-claims',level=2)
add_bullet(doc,'The system does not claim unrestricted general autonomy.')
add_bullet(doc,'A high in-domain approval accuracy does not establish general supervisory competence.')
add_bullet(doc,'Recorded traces do not constitute learning until a validated model or adapter update is trained and promoted.')
add_bullet(doc,'The evaluator/coder agent is a proposal generator, not the authority that certifies or deploys its own changes.')
doc.add_heading('1.4 Operational scope',level=2)
doc.add_paragraph('For this thesis, a small model is an open-weight model with 0.5B-8B parameters that can run in quantized form on one prosumer GPU using no more than 16 GB of inference memory. The funded core strata are 3B-4B and 7B-8B; the 0.5B-1.5B very-small stratum is a conditional extension. A 14B-32B open-weight model and, where affordable, a frontier API model serve as reference ceilings rather than deployment-equivalent baselines.')
doc.add_paragraph('The core study is intentionally bounded to REALM-Bench P6/P9 and J1/J2, plus thesis-specific micro-environments. Adapter routing and broader REALM tiers are conditional extensions after the core architecture and learning claims are supported.')

doc.add_heading('2. Formal System Model',level=1)
doc.add_heading('2.1 Environment',level=2)
doc.add_paragraph('Define the environment as the tuple:')
add_code(doc,'E = <S, O, A, T, C, X, R>')
add_table(doc,['Symbol','Meaning','Required property'],[
    ('S','Canonical world-state space','Typed, versioned, serializable, reconstructable from events.'),
    ('O','Observation function O(s, actor)','Exposes only the state visible and permitted to a given component.'),
    ('A','Action space','Finite typed action schemas plus bounded natural-language arguments.'),
    ('T-hat','Predictive transition model T-hat(s,a)','Estimates expected effects for pre-execution risk checks; it is not the actual external transition.'),
    ('T','Reconciled transition T(s,a,y,e) -> s\'','Applies verified tool outcome y and any exogenous event e after execution.'),
    ('C','Constraint set','Hard invariants, action preconditions, deadlines, dependencies, and policies.'),
    ('X','Exogenous-event space','Disruptions not selected by the agent: delays, cancellations, failures.'),
    ('R','Outcome vector','Safety, protected regression, calibration, effectiveness, recovery, and cost signals with lexicographic promotion priority.'),
],[900,3000,5460])
doc.add_paragraph('The append-only ledger is not identical to the environment. It is an ordered event log L = (e0 ... en) from which the canonical state is reduced: s_n = reduce(s_0, L). A reducer must be deterministic for a fixed schema version.')
doc.add_paragraph('R is not collapsed into an unconstrained reward for promotion. Deployment decisions use the ordering: hard safety, protected-suite regression, calibration, effectiveness, then efficiency. A lower-priority gain cannot compensate for failure of a higher-priority gate.')

doc.add_heading('2.2 Canonical state',level=2)
add_code(doc,'state = {\n  schema_version, episode_id, logical_clock,\n  entities, resources, obligations, dependencies, deadlines,\n  permissions, active_plan, completed_actions, pending_questions,\n  exogenous_events, provenance, revision_chain\n}')
add_table(doc,['State field','Rule'],[
    ('Identity','Every entity, action, plan, obligation, and event has a stable unique ID.'),
    ('Time','Use a single normalized time representation plus explicit timezone and uncertainty.'),
    ('Provenance','Every fact records its source: user, tool, environment, derived rule, or model claim.'),
    ('Epistemic status','Distinguish observed, asserted, inferred, unknown, contradicted, and superseded facts.'),
    ('Revision','Corrections never mutate history; they append a superseding event and recompute active state.'),
    ('Concurrency','Resource reservations and plan versions use optimistic version checks or locks.'),
],[1900,7460])

doc.add_heading('2.3 Action envelope',level=2)
add_code(doc,'action_proposal = {\n  action_id, plan_id, actor_id, action_type, arguments,\n  expected_preconditions, expected_effects, required_capabilities,\n  risk_class, reversibility: REVERSIBLE|COMPENSATABLE|IRREVERSIBLE,\n  compensation_action, postconditions, audit_deadline,\n  evidence_refs, confidence, idempotency_key\n}')
doc.add_paragraph('Workers may create proposals. Only the executor may invoke an external tool, and only after the proposal has passed authorization and validation for the current state version.')

doc.add_heading('2.4 Transition semantics',level=2)
add_number(doc,'Observe: construct actor-specific observation o_t = O(s_t, actor).')
add_number(doc,'Propose: a worker samples a structured proposal a_t from policy pi(a | o_t, goal, adapter).')
add_number(doc,'Authorize: the Supervisor classifies the proposal and recommends approve, repair, query, reject, or escalate.')
add_number(doc,'Validate predictively: deterministic validators check schemas, capabilities, preconditions, preventable invariants, versions, and budgets; T-hat estimates residual risk for effects not fully controlled by the system.')
add_number(doc,'Execute: a narrow executor invokes the permitted tool with an idempotency key.')
add_number(doc,'Reconcile: raw tool output is normalized into a verified outcome event; unverified claimed effects do not become canonical facts.')
add_number(doc,'Transition and audit: the reducer derives s_(t+1); detectable postconditions and invariants are evaluated against actual outcomes.')
add_number(doc,'Commit or recover: the audit resolves to COMMIT, COMPENSATE, REPAIR, or ESCALATE. Irreversible actions require stricter pre-authorization because compensation is unavailable.')
doc.add_paragraph('A COMPENSATABLE proposal is schema-valid only when compensation_action is non-null and has passed pre-validation as an action template. At recovery time, compensation re-enters authorization and validation against the current state, capability, and version; pre-validation does not bypass changed conditions. If compensation is blocked, fails, or cannot be verified by audit_deadline, the system records a detectable failure and escalates. No action may be represented as guaranteed reversible solely because a compensation template exists.')
doc.add_heading('2.5 Constraint enforceability classes',level=2)
add_table(doc,['Class','When evaluated','Example','Failure response'],[
    ('Preventable','Before execution against controlled state','Permission, stale version, known resource capacity.','Block with REPAIR, REJECT, QUERY, or ESCALATE.'),
    ('Predictive','Before execution against T-hat and expected effects','External travel completion or probabilistic tool effect.','Block or escalate when predicted residual risk exceeds policy.'),
    ('Detectable','After reconciliation against actual outcome','Remote service accepted request but produced partial side effect.','Compensate, repair, retry, or escalate.'),
],[1450,2100,3200,2610])

doc.add_heading('3. Runtime Control Plane',level=1)
add_table(doc,['Component','Authority','Forbidden responsibility'],[
    ('Supervisor','Decompose goals, route work, compare proposals, classify risk, request repair or clarification.','Cannot bypass capability checks or directly promote a trained model.'),
    ('Worker','Produce bounded plans, analyses, or typed action proposals for assigned subgoals.','Cannot execute tools outside its granted capability token.'),
    ('Learned verifier','Evaluate semantic ambiguity, soft constraints, and failure likelihood.','Cannot overrule a failed hard invariant.'),
    ('Deterministic validator','Enforce schemas, permissions, preconditions, invariants, budgets, and state versions.','Cannot invent missing semantic facts.'),
    ('Executor','Perform the exact authorized tool call and record raw result.','Cannot alter arguments after approval.'),
    ('Ledger/reducer','Preserve events and derive canonical state.','Cannot accept model prose as an executed event.'),
],[1450,3950,3960])

doc.add_heading('3.1 Decision set and deterministic reconciliation',level=2)
add_table(doc,['Decision','Meaning','Required next step'],[
    ('APPROVE','Proposal is sufficiently specified and passes all required checks.','Execute with current state version and idempotency key.'),
    ('REPAIR','Goal is valid but proposal violates a recoverable constraint.','Return violated constraint IDs and repair target.'),
    ('QUERY','Required state is unknown or materially ambiguous.','Ask the minimal disambiguating question.'),
    ('REJECT','Action is impermissible or irreconcilable with the goal/policy.','Do not execute; provide bounded explanation.'),
    ('ESCALATE','Residual risk or authority exceeds autonomous policy.','Transfer to user or authorized reviewer.'),
],[1300,3960,4100])
doc.add_paragraph('These decisions are typed resolution operations, not a total order. A deterministic policy reconciles the learned recommendation with validator evidence. A learned APPROVE is necessary but never sufficient when a hard check applies.')
add_table(doc,['Supervisor','Validator evidence','Resolved result'],[
    ('APPROVE','HARD + PASS','APPROVE, subject to every remaining gate.'),
    ('APPROVE','HARD + FAIL; repair exists','REPAIR with constraint IDs and repair target.'),
    ('APPROVE','HARD + FAIL; impermissible','REJECT.'),
    ('APPROVE','HARD + UNKNOWN','QUERY when answerable; otherwise ESCALATE.'),
    ('APPROVE','SOFT + FAIL or UNKNOWN','Apply declared residual-risk policy; APPROVE only below the action-specific ceiling.'),
    ('REPAIR','Any','REPAIR unless the evidence requires REJECT or ESCALATE.'),
    ('QUERY','Any','QUERY; no execution while required information is absent.'),
    ('REJECT','Any','REJECT.'),
    ('ESCALATE','Any','ESCALATE.'),
],[1800,3100,4460])
add_code(doc,'permit(a,s) = supervisor_accept(a,s) AND capability_ok(a) AND schema_ok(a)\n  AND preconditions_hold(a,s) AND preventable_invariants_hold(a,s)\n  AND predicted_risk(T_hat(s,a)) <= risk_ceiling(risk_class(a), reversibility(a))\n  AND version_current(a,s) AND budget_ok(a)')
doc.add_paragraph('UNKNOWN is fail-closed for every hard validator: it resolves to QUERY or ESCALATE and can never contribute to execution permission. UNKNOWN on a soft constraint contributes to residual risk under an explicit policy and must remain visible in the trace.')

doc.add_heading('3.2 Hard invariants',level=2)
add_bullet(doc,'No tool executes without an explicit scoped capability grant.')
add_bullet(doc,'No action executes against a stale state or plan version without revalidation.')
add_bullet(doc,'No model-generated claim is recorded as a completed external effect without tool or environment evidence.')
add_bullet(doc,'No irreversible or high-risk action proceeds below its required approval tier.')
add_bullet(doc,'No resource may be double-booked beyond declared capacity.')
add_bullet(doc,'No deadline/dependency violation may be silently converted into a successful outcome.')
add_bullet(doc,'Every promoted model, adapter, prompt, evaluator, schema, and environment has a reproducible version identifier.')
add_bullet(doc,'Every actual tool outcome is audited after reconciliation; failed detectable postconditions trigger compensation, repair, or escalation rather than silent success.')

doc.add_heading('3.3 Revision semantics',level=2)
add_table(doc,['User form','Interpretation','State operation'],[
    ('“Actually, replace X with Y.”','Explicit supersession','Append REVISE event; mark X superseded; recompute dependent plan nodes.'),
    ('“Also add Y.”','Addition','Append new obligation or constraint; preserve X.'),
    ('“Do Y instead.”','Potential replacement','Query if the scope of “instead” is ambiguous.'),
    ('“Cancel that.”','Cancellation','Resolve referent; append CANCEL; run compensation policy for completed effects.'),
],[1800,3000,4560])

doc.add_heading('3.4 Clarification user model',level=2)
doc.add_paragraph('The QUERY pathway requires a versioned user-response environment. The primary scientific condition is a deterministic oracle user that reveals only requested, answerable facts from private state. Naturalistic and adversarial users are robustness conditions, not the ground-truth grader.')
add_code(doc,'user_model = {\n  user_model_id, version, private_world_state_ref, answerable_facts,\n  response_policy, ambiguity_policy, correction_policy, noise_policy,\n  maximum_turns, oracle_access_rules, template_family, seed\n}')
add_bullet(doc,'The acting policy never receives private_world_state directly.')
add_bullet(doc,'The simulator exposes only facts permitted by its response policy and cannot announce global plan correctness.')
add_bullet(doc,'Paired baselines use identical simulator versions, private states, and seeds.')
add_bullet(doc,'Evaluation measures necessary, unnecessary, insufficient, repeated, and incorrectly scoped clarification.')

doc.add_heading('4. Trace and Evidence Contract',level=1)
doc.add_paragraph('The primary learning record is the externally grounded decision trajectory, not unrestricted hidden reasoning. A trace step is:')
add_code(doc,'tau_t = <canonical_state_ref[EVALUATOR_ONLY],\n         actor_observation[POLICY_VISIBLE], goal_ref, policy_version, adapter_version,\n         proposal, supervisor_decision, validator_results, tool_result,\n         transition_diff, outcome_vector, evaluator_labels, provenance>')
doc.add_paragraph('Canonical state is retained for replay and grading, but policy-training inputs must be reconstructed exclusively from actor_observation and other POLICY_VISIBLE fields. The data compiler rejects examples that expose PRIVATE_ENVIRONMENT or EVALUATOR_ONLY fields to the deployed policy.')
add_table(doc,['Access class','Permitted consumer'],[
    ('POLICY_VISIBLE','The worker or policy being trained and deployed.'),
    ('SUPERVISOR_VISIBLE','Supervisor-only coordination and risk context.'),
    ('VALIDATOR_VISIBLE','Deterministic validation and capability enforcement.'),
    ('EVALUATOR_ONLY','Replay, grading, outcome attribution, and research analysis.'),
    ('PRIVATE_ENVIRONMENT','Hidden world state and user-simulator facts; never a policy input.'),
],[2200,7160])
add_table(doc,['Trace class','Retain','Use'],[
    ('Successful','Minimal valid plan, actions, outcomes, costs.','Behavior cloning, positive preferences, efficiency optimization.'),
    ('Recoverable failure','Bad proposal, violated constraint, successful repair.','Contrastive pairs and repair-policy training.'),
    ('Unsafe/impermissible','Proposal plus deterministic violation evidence.','Supervisor and verifier negatives; adversarial regression tests.'),
    ('Ambiguous','Missing facts, question asked, user resolution.','Clarification policy and uncertainty calibration.'),
    ('Tool/environment failure','Authorized action, tool error, recovery sequence.','Retry, compensation, and robust replanning.'),
],[1800,3700,3860])
doc.add_heading('4.1 Data quality gates',level=2)
add_bullet(doc,'Redact secrets and private data before a trace enters the training store.')
add_bullet(doc,'Canonicalize tool names, action schemas, state serialization, and decision labels.')
add_bullet(doc,'Deduplicate near-identical trajectories before splitting.')
add_bullet(doc,'Split by scenario template, generator seed, and entity graph - not merely by individual trace.')
add_bullet(doc,'Keep benchmark test instances and regression-only cases out of training selection.')
add_bullet(doc,'Record evaluator disagreement and abstention rather than forcing every trace into a label.')
add_bullet(doc,'Prefer executable outcomes and formal constraint checks over model-only judgments.')

doc.add_heading('5. Optimization and Post-Training Plane',level=1)
doc.add_heading('5.1 Definition of an improvement cycle',level=2)
add_code(doc,'Cycle k: freeze M_k -> collect D_k -> grade -> select -> train candidate M_(k+1)\n         -> evaluate frozen suites -> safety gate -> promote or quarantine -> retain rollback')
add_number(doc,'Freeze the deployed policy, prompts, adapters, evaluator versions, and environment schema.')
add_number(doc,'Collect traces under declared exploration and inference budgets.')
add_number(doc,'Grade with deterministic outcomes first, independent learned evaluators second, and human adjudication for disputed high-impact cases.')
add_number(doc,'Select data using failure taxonomy, novelty, confidence, balance, and contamination controls.')
add_number(doc,'Train one isolated intervention: prompt, SFT, DPO, process reward, agent RL, or adapter.')
add_number(doc,'Evaluate against frozen in-domain, held-out, adversarial, transfer, calibration, and regression suites.')
add_number(doc,'Promote only if preregistered thresholds are met; otherwise quarantine artifacts and log diagnosis.')

doc.add_heading('5.2 Training-method decision rules',level=2)
add_table(doc,['Method','Use when','Do not infer'],[
    ('SFT / behavior cloning','There are verified target actions or repaired trajectories.','Imitation alone learns long-horizon recovery or calibrated refusal.'),
    ('DPO','Pairs differ in plan feasibility, repair quality, safety, or efficiency.','Offline DPO alone closes the full self-improvement loop.'),
    ('Process reward model','Intermediate actions need progress/promise estimates and terminal reward is sparse.','A learned process score is a hard safety guarantee.'),
    ('Offline agent RL','Logged transitions have defensible rewards and adequate behavior coverage.','Off-policy estimates are reliable outside trace support.'),
    ('On-policy agent RL','A sandboxed environment supplies executable rewards and safe exploration.','Reward optimization cannot exploit evaluator blind spots.'),
],[1650,3910,3800])

doc.add_heading('5.3 Promotion gate',level=2)
callout(doc,'Promotion rule','Promotion is lexicographic, not a weighted sum: pass hard safety first, protected regression second, calibration third, effectiveness fourth, and efficiency fifth. A lower-priority improvement cannot offset a failed higher-priority gate. Evaluation artifacts must be produced by a pipeline the candidate could not modify.')
add_code(doc,'promote(c) iff unsafe_rate_upper_bound_95(c) <= epsilon_safety\n  AND protected_regression(c) <= epsilon_regression\n  AND ECE(c) <= epsilon_calibration\n  AND primary_delta(c) > delta_min with CI requirement\n  AND cost(c) <= deployment_budget')

doc.add_heading('6. Modular LoRA Skill Layer',level=1)
doc.add_paragraph('Adapters are versioned policy modules, not informal “power boosts.” Each adapter has a manifest:')
add_code(doc,'adapter_manifest = {adapter_id, base_model_hash, training_data_version, domain,\n  capability_scope, rank, target_modules, router_features, eval_report,\n  known_failures, incompatibilities, fallback_adapter, status}')
add_table(doc,['Stage','Research rule'],[
    ('General adapter','First distill a general action-selection and tool-use policy.'),
    ('Specialists','Add only domains with distinct trace distributions and measurable oracle-routing gains.'),
    ('Oracle routing','Evaluate the upper bound using ground-truth task/domain labels.'),
    ('Learned routing','Train a lightweight router only if oracle routing improves the Pareto frontier.'),
    ('Composition','Begin with one active specialist. Treat merging or simultaneous activation as a separate interference experiment.'),
    ('Fallback','Low confidence, incompatible version, or failed health check routes to the general adapter.'),
],[1750,7610])
doc.add_heading('6.1 Adapter routing experiment',level=2)
add_bullet(doc,'No adapter: frozen base model.')
add_bullet(doc,'General adapter only.')
add_bullet(doc,'Oracle-selected specialist.')
add_bullet(doc,'Learned router with confidence threshold.')
add_bullet(doc,'Random and deliberately wrong adapter controls.')
add_bullet(doc,'Optional composition experiment after single-adapter routing is understood.')

doc.add_heading('7. Research Questions and Hypotheses',level=1)
add_table(doc,['ID','Question / hypothesis','Primary endpoint'],[
    ('H1','Adding a deterministic validator to supervisor-worker orchestration reduces hard-constraint violations without an unacceptable loss in task completion.','Violation rate and completion delta.'),
    ('H2','Trace SFT improves small-model planning over architecture-only baselines on held-out scenario templates.','Held-out task success.'),
    ('H3','Preference or process supervision improves repair and clarification behavior beyond trace SFT.','Recovery success and unnecessary-query rate.'),
    ('H4 - conditional','If the core B0-B5 results pass their gates, specialist LoRAs improve the success/compute Pareto frontier when correctly routed.','Success per token/latency and oracle-router gap.'),
    ('H5','Across preregistered versioned cycles, candidates admitted by the promotion gate improve the primary endpoint while remaining inside protected-regression and unsafe-action bounds.','Per-cycle gate pass, primary delta, protected regression.'),
    ('H6','Hybrid learned/deterministic gating is better calibrated and safer than learned supervision alone.','Unsafe FN rate, ECE, selective risk.'),
],[700,5780,2880])

doc.add_heading('8. REALM-Bench Evaluation Protocol',level=1)
doc.add_heading('8.1 Benchmark role',level=2)
doc.add_paragraph('REALM-Bench is the primary planning and disruption testbed, not the sole evidence for general agency. Use its static, dynamic, routing, scheduling, coordination, and JSSP tasks alongside architecture-specific safety and revision suites.')
doc.add_paragraph('The Thanksgiving micro-environment is a formally encoded derivative of REALM-Bench P6/P9, extended by this thesis with executable transitions, revision semantics, clarification interactions, capability policies, counterfactual safety cases, and post-execution compensation. Every result must distinguish original benchmark content from thesis-specific extensions.')
doc.add_heading('8.2 Recommended progression',level=2)
add_table(doc,['Tier','Tasks','Purpose'],[
    ('0 - harness checks','Hand-authored micro-environments','Validate state reducers, invariants, permissions, revisions, idempotency, and trace capture.'),
    ('1 - static','CORE: P6, J1. EXTENSION: P1, P3, P5.','Planning feasibility and baseline optimality.'),
    ('2 - multi-party','EXTENSION: P2 and larger static variants.','Coordination, resource conflicts, and communication overhead.'),
    ('3 - disruption','CORE: P9, J2. EXTENSION: P4, P8.','Reactive replanning, recovery latency, and plan stability.'),
    ('4 - scale','EXTENSION: P7, P10, J3, J4.','Long horizons, dependency depth, resource pressure, and computational efficiency.'),
],[1300,2400,5660])
doc.add_heading('8.3 Split policy',level=2)
add_bullet(doc,'Training: generated instances and trace variants from declared scenario families and seeds.')
add_bullet(doc,'Development: held-out seeds plus controlled counterfactual pairs used for model selection.')
add_bullet(doc,'Test: held-out templates, graph structures, disruption combinations, and scales; never used for data selection.')
add_bullet(doc,'Transfer: entire task families excluded from training, such as train on event logistics and test on ride-sharing or JSSP.')
add_bullet(doc,'Protected regression: historical failures retained for evaluation only after the corresponding training round.')
doc.add_heading('8.4 Versioned instance generation',level=2)
doc.add_paragraph('REALM-Bench provides specifications and partial generation utilities, but this research requires a controlled generator for counterfactuals, held-out graphs, disruptions, and user interactions. Generator design is therefore an experimental variable and part of dataset lineage.')
add_code(doc,'instance_manifest = {generator_id, generator_version, scenario_family, seed,\n  difficulty_parameters, constraint_graph_hash, entity_graph_hash,\n  disruption_schedule_hash, user_model_version, source_dataset_lineage,\n  feasibility_oracle_version, split_assignment\n}')
add_bullet(doc,'Validate generated feasibility with executable checks or solver-backed oracles where possible.')
add_bullet(doc,'Hold out parameter combinations, graph structures, templates, and disruption compositions - not only random seeds.')
add_bullet(doc,'Report original REALM instances separately from generated extensions to expose generator exploitation.')
doc.add_heading('8.5 Metrics',level=2)
add_table(doc,['Category','Metrics'],[
    ('Effectiveness','Task success, goal completeness, feasible-plan rate, makespan/optimality gap.'),
    ('Safety','Hard violations, unsafe false negatives, unnecessary rejection, escalation correctness.'),
    ('Recovery','Disruption recovery rate, repair steps, recovery latency, preserved-plan fraction.'),
    ('Calibration','Expected calibration error, Brier score, risk-coverage curve, abstention utility.'),
    ('Efficiency','Wall time, tokens, tool calls, retries, peak memory, energy if measurable.'),
    ('Learning','Gain per trace, gain per training FLOP, cycle-wise improvement, regression rate, forgetting.'),
    ('Coordination','Conflicting assignments, duplicate work, communication volume, deadlocks.'),
],[1700,7660])
doc.add_heading('8.6 Statistical protocol',level=2)
add_bullet(doc,'Predeclare primary endpoint, safety ceiling, minimum meaningful effect, and protected suites.')
add_bullet(doc,'Use identical instances, tool conditions, and inference budgets for paired architecture comparisons.')
add_bullet(doc,'Run multiple seeds or repeated stochastic trials; report mean, dispersion, and confidence intervals.')
add_bullet(doc,'Use paired bootstrap intervals for success/cost deltas and exact or Wilson intervals for rare safety events.')
add_bullet(doc,'Power safety gates separately from effectiveness tests. For a 1% unsafe-rate ceiling with zero observed violations, use at least 300 independent safety trials per evaluated candidate to obtain an approximately 1% one-sided 95% exact-binomial upper bound; increase the sample when violations occur or when dependence reduces effective sample size.')
add_bullet(doc,'Preregister epsilon_safety and derive the required safety-suite size from an exact binomial calculation before evaluation. Do not infer a 1% ceiling from only 100 trials; zero violations in 100 still permits a true rate of roughly 3% at a one-sided 95% bound.')
add_bullet(doc,'Correct for multiple comparisons when testing many adapters, models, or task subsets.')
add_bullet(doc,'Report both aggregate results and per-task failure distributions; avoid hiding catastrophic tails in means.')

doc.add_heading('9. Baselines and Ablations',level=1)
doc.add_paragraph('Core thesis matrix: P6, P9, J1, and J2 with B0-B5. B6-B8 and broader REALM tiers are conditional extensions rather than required core experiments.')
add_table(doc,['ID','Configuration','Isolated question'],[
    ('B0','Prompted small model, no tools','What can the base model plan directly?'),
    ('B1','Single looping agent with tools','What does environment interaction add?'),
    ('B2','Supervisor + workers','What does role decomposition add?'),
    ('B3','B2 + deterministic validator','What does the safety shield add?'),
    ('B4','B3 + trace SFT','What does distilled behavior add?'),
    ('B5','B4 + DPO or process supervision','What does preference/process learning add?'),
    ('B6','B5 + oracle specialist LoRA','Is specialization potentially useful?'),
    ('B7','B5 + learned router','How much of oracle benefit survives routing?'),
    ('B8','Full versioned optimization loop','Does repeated improvement outperform one-shot training?'),
],[700,3650,5010])
doc.add_heading('9.1 Controls',level=2)
add_bullet(doc,'Hold base model, quantization, context limit, tool set, state schema, prompts, temperature, and token budget constant where the comparison permits.')
add_bullet(doc,'Compare against a classical solver or executable oracle for constraint problems when available.')
add_bullet(doc,'Include equal-compute and equal-latency comparisons, not only equal model size.')
add_bullet(doc,'Include a larger-model teacher/reference to measure the remaining capability gap, not as the only baseline.')
doc.add_heading('9.2 Feasibility budget and stopping rules',level=2)
add_table(doc,['Resource','Core-study ceiling'],[
    ('Models','One primary 3B-4B student, one 7B-8B upper-small reference, and one larger teacher/reference.'),
    ('Core tasks','P6, P9, J1, J2 plus micro-tests; broader tiers are stretch work.'),
    ('Baselines','B0-B5 required; B6-B8 conditional on successful core results.'),
    ('Training traces','Pilot 1,000-3,000; core ceiling 20,000 verified steps unless preregistered otherwise.'),
    ('Effectiveness evaluation','Minimum 100 paired instances per core family where generation permits, with repeated stochastic trials.'),
    ('Safety evaluation','Separately powered from effectiveness: at least 300 independent trials per candidate for a preregistered 1% ceiling when zero violations are required; recompute n for the chosen ceiling and dependence structure.'),
    ('Improvement cycles','Three as a repeatability demonstration; no long-term monotonic-trend claim.'),
    ('Compute','Record and cap GPU-hours before each stage; stop methods that exceed budget without development-set benefit.'),
    ('Stopping','Stop a branch for safety-gate failure, protected regression, no meaningful development gain, or exhausted compute budget.'),
],[1800,7560])

doc.add_heading('10. Implementation Interfaces',level=1)
doc.add_heading('10.1 Validator result',level=2)
add_code(doc,'validator_result = {\n  validator_id, validator_version, status: PASS|FAIL|UNKNOWN,\n  criticality: HARD|SOFT, enforceability: PREVENTABLE|PREDICTIVE|DETECTABLE,\n  constraint_ids, evidence_refs, observed_state_version, uncertainty_reason,\n  permitted_resolution: EXECUTE|REPAIR|QUERY|REJECT|ESCALATE,\n  explanation_code, repair_hints, timestamp\n}')
doc.add_heading('10.2 Supervisor decision',level=2)
add_code(doc,'supervisor_decision = {\n  decision: APPROVE|REPAIR|QUERY|REJECT|ESCALATE,\n  confidence, violated_constraint_ids, uncertainty_reasons,\n  requested_information, proposed_repair, evidence_refs,\n  supervisor_model_version, adapter_version\n}')
doc.add_heading('10.3 Outcome vector',level=2)
add_code(doc,'outcome = {\n  goal_success, constraint_score, hard_violation_count,\n  optimality_gap, recovery_success, execution_cost, latency,\n  token_count, tool_calls, user_interventions,\n  clarification_total, clarification_necessary, clarification_unnecessary,\n  clarification_insufficient, clarification_repeated,\n  clarification_incorrect_scope, evaluator_disagreement\n}')
doc.add_heading('10.4 Minimum artifact registry',level=2)
add_bullet(doc,'Model and quantization manifests.')
add_bullet(doc,'Prompt and graph definitions.')
add_bullet(doc,'Environment schema and reducer versions.')
add_bullet(doc,'Tool schemas and capability policies.')
add_bullet(doc,'Dataset lineage and split manifests.')
add_bullet(doc,'Instance-generator and user-simulator manifests, versions, templates, and seeds.')
add_bullet(doc,'Evaluator, reward, and validator versions.')
add_bullet(doc,'Training configuration, seeds, checkpoints, and adapter manifests.')
add_bullet(doc,'Evaluation reports, promotion decisions, and rollback pointers.')

doc.add_heading('11. Failure Taxonomy',level=1)
add_table(doc,['Layer','Failure examples'],[
    ('Observation','Missing fact, stale state, excess context, hidden dependency, incorrect provenance.'),
    ('Planning','Infeasible ordering, omitted obligation, resource conflict, poor decomposition, horizon collapse.'),
    ('Coordination','Duplicate work, inconsistent subplans, deadlock, supervisor bottleneck, message loss.'),
    ('Authorization','Over-approval, over-refusal, authority confusion, confidence miscalibration.'),
    ('Validation','Incomplete rules, wrong constraint encoding, UNKNOWN treated as PASS, version mismatch.'),
    ('Execution','Wrong arguments, non-idempotent retry, partial side effect, timeout, tool-schema drift.'),
    ('Learning','Label leakage, evaluator circularity, reward hacking, contamination, forgetting, adapter interference.'),
    ('Evaluation','Weak oracle, benchmark memorization, metric gaming, insufficient seeds, hidden cost trade-off.'),
],[1500,7860])

doc.add_heading('12. Staged Research Plan',level=1)
add_table(doc,['Stage','Deliverable','Exit criterion'],[
    ('1. Formal environment','Typed schema, reducer, predictive/actual transitions, constraint classes, revision, compensation, user-model, and generator contracts.','Deterministic replay, leakage, UNKNOWN, and invariant tests pass.'),
    ('2. Baselines','B0-B3 on micro-environments and initial REALM tasks.','Reproducible paired results and trace completeness.'),
    ('3. Evidence pipeline','Trace store, graders, failure taxonomy, data lineage, frozen splits.','Independent rerun reproduces scores and datasets.'),
    ('4. General distillation','Worker SFT plus repaired-failure contrasts.','Held-out gain over B3 at bounded safety regression.'),
    ('5. Supervisor learning','SFT/classifier, DPO, and process-supervision comparison.','Improved selective risk and recovery, not accuracy alone.'),
    ('6. Adapter study - conditional','General/specialist LoRAs, oracle and learned routing.','Begin only after B0-B5 support the core thesis; require Pareto improvement.'),
    ('7. Repeated cycles','Three frozen collect-train-evaluate-promote cycles as a repeatability demonstration.','Admitted candidates satisfy every lexicographic gate; no monotonic-trend claim.'),
    ('8. Stress and release','Transfer, adversarial, evaluator-failure, rollback, and reproducibility study.','Claims supported by ablations, intervals, artifacts, and limitations.'),
],[1450,3950,3960])

doc.add_heading('13. Immediate Build Backlog',level=1)
for item in [
    'Write JSON Schemas for canonical state, access-scoped observations, action proposal, validator result, supervisor decision, tool outcome, trace step, episode, user model, instance manifest, and adapter manifest.',
    'Implement a pure deterministic state reducer and event-replay test harness.',
    'Encode the P6/P9-derived Thanksgiving micro-environment as explicit resources, obligations, time windows, predictive effects, actual outcomes, constraint classes, and compensation rules.',
    'Create counterfactual state pairs that differ in one safety-critical fact.',
    'Separate Supervisor recommendation from deterministic execution permission in the orchestration graph.',
    'Add version IDs and provenance to every trace artifact.',
    'Build the deterministic user simulator and clarification metrics before training the QUERY policy.',
    'Build and version the controlled instance generator; record graph hashes, disruption schedules, lineage, and solver validation.',
    'Add a data-compiler test that rejects EVALUATOR_ONLY and PRIVATE_ENVIRONMENT fields from policy inputs.',
    'Build B0-B3 runners with equal budgets and paired seeds.',
    'Freeze the core REALM subset: P6, P9, J1, and J2. Treat P1 and broader tiers as extensions.',
    'Create a classical feasibility checker or solver-backed oracle for the selected tasks.',
    'Freeze a first evaluation manifest before collecting training traces.'
]: add_number(doc,item)

doc.add_heading('14. Success Criteria for the Thesis',level=1)
callout(doc,'Minimum publishable result','On the scoped P6/P9/J1/J2 study, the architecture shows a statistically supported improvement over single-agent and learned-only baselines; hard violations remain below a preregistered ceiling; gains survive equal-compute comparisons; ablations identify the causal components; and three versioned cycles demonstrate repeatable gate-controlled improvement without claiming a long-term monotonic trend.')
doc.add_paragraph('A stronger result would show that a routed small-model system approaches or exceeds a substantially larger model on the success-safety-cost Pareto frontier, particularly under disruptions and revision semantics.')

doc.add_heading('15. Selected Research Foundations',level=1)
sources=[
('REALM-Bench: planning, coordination, disruption, and scheduling evaluation','https://arxiv.org/abs/2502.18836'),
('REALM-Bench repository and datasets','https://github.com/genglongling/REALM-Bench'),
('Direct Preference Optimization','https://arxiv.org/abs/2305.18290'),
('Online preference optimization distinctions','https://arxiv.org/abs/2403.08635'),
('LoRA: Low-Rank Adaptation of Large Language Models','https://arxiv.org/abs/2106.09685'),
('Agent distillation into small models with retrieval and code tools','https://arxiv.org/abs/2505.17612'),
('Process reward models for LLM agents','https://arxiv.org/abs/2502.10325'),
('AgentPRM: step-wise promise and progress','https://arxiv.org/abs/2511.08325'),
('ReAct: reasoning and acting in language models','https://arxiv.org/abs/2210.03629'),
('DSPy: compiling declarative LM pipelines','https://arxiv.org/abs/2310.03714'),
]
for title,url in sources:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(title+': ').bold=True; p.add_run(url)

doc.add_heading('Appendix A. Formal Separation of Authorities',level=1)
add_table(doc,['Artifact / operation','May propose','May validate','May execute/promote'],[
    ('Tool action','Worker or Supervisor','Learned verifier + deterministic validator','Executor only'),
    ('Constraint change','Evaluator/coder or researcher','Schema tests + authorized review','Configuration authority only'),
    ('Training example','Selector agent or pipeline','Data-quality gates + lineage checks','Training pipeline only'),
    ('Model update','Trainer/coder agent','Frozen evaluation pipeline','Promotion authority only'),
    ('Adapter route','Router','Compatibility and confidence gates','Runtime adapter manager'),
    ('Evaluator update','Coder agent or researcher','Meta-evaluation and frozen canaries','Independent evaluation authority'),
    ('Compensation action','Original planner or recovery controller','Pre-validate template; re-authorize and revalidate current state at recovery time','Executor only; blocked or failed compensation escalates'),
],[2050,2350,2510,2450])

doc.add_heading('Appendix B. Terminology',level=1)
add_table(doc,['Term','Operational definition'],[
    ('Agent','A model policy plus harness, observation interface, memory/state access, tool schemas, permissions, and an iterative control loop.'),
    ('Environment','A state-transition system that supplies observations, accepts permitted actions, produces outcomes, and may generate disruptions.'),
    ('Action space','The typed set of proposals the agent may emit; executable actions are the authorized subset at a particular state.'),
    ('Harness','The orchestration, prompting, context construction, parsing, retries, permissions, validators, tools, and logging around a model.'),
    ('Self-improvement','A measured sequence of versioned policy or harness updates derived from prior traces and admitted through independent promotion gates.'),
    ('Post-training','Any parameter update after base pretraining, including SFT, preference optimization, reward-model training, RL, and LoRA adaptation.'),
    ('Conditional autonomy','Authority to act only when state, confidence, risk, permissions, and constraints satisfy an explicit execution policy.'),
    ('Small model','An open-weight 0.5B-8B parameter model that can run quantized on one prosumer GPU within 16 GB of inference memory.'),
    ('Thanksgiving micro-environment','A P6/P9-derived thesis environment with executable transitions, clarification, revision, safety counterfactuals, and compensation semantics.'),
],[1900,7460])

doc.core_properties.title='Conditional Autonomy for Self-Improving Small-Model Agents'
doc.core_properties.subject='Formal thesis architecture and research protocol'
doc.core_properties.author='Research Architecture Working Draft'
doc.save(OUT)
print(OUT.resolve())
